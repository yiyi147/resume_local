from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_word(data: dict, template_name: str = "modern") -> bytes:
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    if data.get('basics'):
        basics = data['basics']
        p = doc.add_heading(basics.get('name', ''), 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact = []
        if basics.get('email'):
            contact.append(basics['email'])
        if basics.get('phone'):
            contact.append(basics['phone'])
        if basics.get('location'):
            contact.append(basics['location'])
        if contact:
            p = doc.add_paragraph(' | '.join(contact))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if basics.get('summary'):
            doc.add_paragraph(basics['summary'])
    
    if data.get('education'):
        doc.add_heading('教育背景', level=1)
        for edu in data['education']:
            p = doc.add_paragraph()
            run = p.add_run(edu.get('school', ''))
            run.bold = True
            if edu.get('date'):
                p.add_run(f" ({edu['date']})")
            if edu.get('degree') or edu.get('major'):
                doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('major', '')}")
            if edu.get('description'):
                doc.add_paragraph(edu['description'])
    
    if data.get('work'):
        doc.add_heading('工作经历', level=1)
        for work in data['work']:
            p = doc.add_paragraph()
            run = p.add_run(work.get('company', ''))
            run.bold = True
            if work.get('date'):
                p.add_run(f" ({work['date']})")
            if work.get('position'):
                doc.add_paragraph(work['position'])
            if work.get('description'):
                doc.add_paragraph(work['description'])
            if work.get('highlights'):
                for h in work['highlights']:
                    doc.add_paragraph(h, style='List Bullet')
    
    if data.get('projects'):
        doc.add_heading('项目经历', level=1)
        for proj in data['projects']:
            p = doc.add_paragraph()
            run = p.add_run(proj.get('name', ''))
            run.bold = True
            if proj.get('date'):
                p.add_run(f" ({proj['date']})")
            if proj.get('description'):
                doc.add_paragraph(proj['description'])
            if proj.get('highlights'):
                for h in proj['highlights']:
                    doc.add_paragraph(h, style='List Bullet')
    
    if data.get('skills'):
        doc.add_heading('专业技能', level=1)
        skills_text = ', '.join([s.get('name', s) if isinstance(s, dict) else s for s in data['skills']])
        doc.add_paragraph(skills_text)
    
    if data.get('awards'):
        doc.add_heading('荣誉奖项', level=1)
        for award in data['awards']:
            p = doc.add_paragraph()
            run = p.add_run(award.get('title', ''))
            run.bold = True
            if award.get('date'):
                p.add_run(f" ({award['date']})")
            if award.get('description'):
                doc.add_paragraph(award['description'])
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

import io
